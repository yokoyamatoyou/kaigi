import logging
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

import docx
import mammoth
import PyPDF2

from .api_clients import BaseAIClient
from .models import AppConfig, DocumentSummary, FileInfo  # AppConfig をインポート
from .utils import Timer, chunk_text, count_tokens, extract_content_and_tokens

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """テキスト抽出結果"""
    extracted_text: str
    metadata: Dict[str, Any]
    error_message: Optional[str] = None
    
    @property
    def is_success(self) -> bool:
        """抽出が成功したかどうか"""
        return self.error_message is None


class DocumentProcessor:
    """ドキュメント処理クラス"""
    
    def __init__(self, config: AppConfig): # AppConfig を受け取るように変更
        """
        初期化
        
        Args:
            config: アプリケーション設定
        """
        self.config = config # config をインスタンス変数として保持
        self.max_file_size_bytes = self.config.max_document_size_mb * 1024 * 1024
        self.supported_extensions = {'.docx', '.pdf', '.txt'} # txtも追加
        
        logger.info(f"DocumentProcessor 初期化完了: 最大ファイルサイズ={self.config.max_document_size_mb}MB")
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """
        ファイルの妥当性をチェック
        
        Args:
            file_path: ファイルパス
        
        Returns:
            (is_valid, error_message): 妥当性チェック結果とエラーメッセージ
        """
        path = Path(file_path)
        
        if not path.exists():
            return False, f"ファイルが存在しません: {file_path}"
        
        file_size = path.stat().st_size
        if file_size > self.max_file_size_bytes:
            size_mb = file_size / (1024 * 1024)
            return False, f"ファイルサイズが上限を超えています: {size_mb:.1f}MB > {self.config.max_document_size_mb}MB"
        
        if path.suffix.lower() not in self.supported_extensions:
            return False, f"サポートされていないファイル形式: {path.suffix}"
        
        try:
            with open(file_path, 'rb') as f:
                f.read(1)
        except PermissionError:
            return False, f"ファイルの読み取り権限がありません: {file_path}"
        except Exception as e:
            return False, f"ファイル読み取りエラー: {str(e)}"
        
        return True, ""
    
    def extract_text_from_docx(self, file_path: str, use_mammoth: bool = True) -> ExtractionResult:
        """
        DOCXファイルからテキストを抽出
        
        Args:
            file_path: DOCXファイルのパス
            use_mammoth: mammothライブラリを使用するかどうか（リッチテキスト対応）
        
        Returns:
            ExtractionResult: 抽出結果
        """
        try:
            with Timer(f"DOCX抽出 ({Path(file_path).name})"):
                file_info = self._get_file_info(file_path)
                
                extracted_text = ""
                metadata = {
                    "file_path": file_path,
                    "file_size": file_info.size_bytes,
                    "extraction_method": "",
                    "extraction_timestamp": datetime.now(),
                    "document_properties": {}
                }
                
                if use_mammoth:
                    try:
                        with open(file_path, "rb") as docx_file:
                            result = mammoth.extract_raw_text(docx_file)
                            extracted_text = result.value
                            metadata["extraction_method"] = "mammoth"
                            if result.messages:
                                logger.warning(f"Mammoth警告: {[msg.message for msg in result.messages]}")
                    except Exception as e:
                        logger.warning(f"Mammothによる抽出失敗、python-docxに切り替え: {e}")
                        use_mammoth = False
                
                if not use_mammoth or not extracted_text.strip():
                    doc = docx.Document(file_path)
                    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                paragraphs.append(" | ".join(row_text))
                    extracted_text = "\n\n".join(paragraphs)
                    metadata["extraction_method"] = "python-docx"
                    try:
                        metadata["document_properties"] = {
                            "title": doc.core_properties.title or "", "author": doc.core_properties.author or "",
                            "subject": doc.core_properties.subject or "",
                            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                        }
                    except Exception as e: logger.warning(f"ドキュメントプロパティ取得失敗: {e}")
                
                extracted_text = self._clean_extracted_text(extracted_text)
                metadata.update({
                    "text_length": len(extracted_text),
                    "paragraph_count": len(extracted_text.split('\n\n')),
                    "word_count": len(extracted_text.split()) if extracted_text else 0
                })
                logger.info(f"DOCX抽出完了: {len(extracted_text)}文字, {metadata['word_count']}語")
                return ExtractionResult(extracted_text=extracted_text, metadata=metadata)
                
        except Exception as e:
            error_msg = f"DOCX抽出エラー: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return ExtractionResult(extracted_text="", metadata={"file_path": file_path}, error_message=error_msg)
    
    def extract_text_from_pdf(self, file_path: str) -> ExtractionResult:
        """
        PDFファイルからテキストを抽出
        """
        try:
            with Timer(f"PDF抽出 ({Path(file_path).name})"):
                file_info = self._get_file_info(file_path)
                extracted_text = ""
                metadata = {
                    "file_path": file_path, "file_size": file_info.size_bytes,
                    "extraction_method": "PyPDF2", "extraction_timestamp": datetime.now(), "pdf_info": {}
                }
                
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    try:
                        pdf_info_meta = pdf_reader.metadata
                        if pdf_info_meta:
                            metadata["pdf_info"] = {
                                k.lstrip('/'): str(v) for k, v in pdf_info_meta.items()
                                if k in ['/Title', '/Author', '/Subject', '/Creator', '/Producer', '/CreationDate', '/ModDate']
                            }
                    except Exception as e: logger.warning(f"PDF情報取得失敗: {e}")
                    
                    num_pages = len(pdf_reader.pages)
                    metadata["page_count"] = num_pages
                    pages_text = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                pages_text.append(f"[ページ {page_num + 1}]\n{page_text.strip()}")
                        except Exception as e:
                            logger.warning(f"ページ {page_num + 1} 抽出失敗: {e}")
                            pages_text.append(f"[ページ {page_num + 1}]\n（抽出失敗）")
                    extracted_text = "\n\n".join(pages_text)
                
                extracted_text = self._clean_extracted_text(extracted_text)
                metadata.update({
                    "text_length": len(extracted_text),
                    "word_count": len(extracted_text.split()) if extracted_text else 0
                })
                logger.info(f"PDF抽出完了: {num_pages}ページ, {len(extracted_text)}文字, {metadata['word_count']}語")
                return ExtractionResult(extracted_text=extracted_text, metadata=metadata)
                
        except Exception as e:
            error_msg = f"PDF抽出エラー: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return ExtractionResult(extracted_text="", metadata={"file_path": file_path}, error_message=error_msg)

    def extract_text_from_txt(self, file_path: str) -> ExtractionResult:
        """TXTファイルからテキストを抽出"""
        try:
            with Timer(f"TXT抽出 ({Path(file_path).name})"):
                file_info = self._get_file_info(file_path)
                metadata = {
                    "file_path": file_path, "file_size": file_info.size_bytes,
                    "extraction_method": "direct_read", "extraction_timestamp": datetime.now()
                }
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:  # UTF-8をデフォルトに
                        extracted_text = f.read()
                except UnicodeDecodeError:
                    logger.warning(f"UTF-8 decoding failed for {file_path}, attempting fallback encoding")
                    try:
                        import chardet  # type: ignore
                        with open(file_path, 'rb') as fb:
                            raw = fb.read()
                        detected = chardet.detect(raw)
                        fallback_encoding = detected.get('encoding') or 'utf-8'
                        with open(file_path, 'r', encoding=fallback_encoding, errors='ignore') as f:
                            extracted_text = f.read()
                        logger.info(f"Fallback encoding used: {fallback_encoding}")
                    except Exception:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            extracted_text = f.read()
                        logger.info("Fallback encoding used: utf-8 with errors='ignore'")
                
                extracted_text = self._clean_extracted_text(extracted_text) # クリーニングは適用
                metadata.update({
                    "text_length": len(extracted_text),
                    "word_count": len(extracted_text.split()) if extracted_text else 0
                })
                logger.info(f"TXT抽出完了: {len(extracted_text)}文字, {metadata['word_count']}語")
                return ExtractionResult(extracted_text=extracted_text, metadata=metadata)

        except Exception as e:
            error_msg = f"TXT抽出エラー: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return ExtractionResult(extracted_text="", metadata={"file_path": file_path}, error_message=error_msg)

    def extract_text(self, file_path: str) -> ExtractionResult:
        is_valid, error_message = self.validate_file(file_path)
        if not is_valid:
            return ExtractionResult(extracted_text="", metadata={"file_path": file_path}, error_message=error_message)
        
        extension = Path(file_path).suffix.lower()
        if extension == '.docx': return self.extract_text_from_docx(file_path)
        elif extension == '.pdf': return self.extract_text_from_pdf(file_path)
        elif extension == '.txt': return self.extract_text_from_txt(file_path)
        else:
            return ExtractionResult(extracted_text="", metadata={"file_path": file_path}, error_message=f"サポートされていないファイル形式: {extension}")
    
    def _clean_extracted_text(self, text: str) -> str:
        if not text: return ""
        text = re.sub(r'[ \t]+', ' ', text) # 連続する空白・タブを単一の空白に
        text = re.sub(r'\n\s*\n', '\n\n', text) # 複数の空行を1つの空行に
        text = re.sub(r'\n{3,}', '\n\n', text) # 3つ以上の連続改行を2つに
        text = '\n'.join(line.rstrip() for line in text.split('\n'))
        return text.strip()
    
    def _get_file_info(self, file_path: str) -> FileInfo:
        path = Path(file_path)
        file_type_str = path.suffix.lower().lstrip('.')
        # FileInfoのfile_typeはLiteralなので、対応する値にキャストする
        # 簡単のため、ここではサポートされている拡張子のみを想定
        if file_type_str not in ["docx", "pdf", "txt"]: 
            # この状況はvalidate_fileで弾かれるはずだが、念のため
            logger.warning(f"予期しないファイルタイプ: {file_type_str}。'txt'として処理します。")
            file_type_literal = "txt" 
        else:
            file_type_literal = file_type_str # type: ignore

        return FileInfo(
            filename=path.name, filepath=str(path.absolute()),
            file_type=file_type_literal, size_bytes=path.stat().st_size
        )

    async def summarize_document_for_meeting(
        self,
        text: str,
        summarizer_ai_client: BaseAIClient,
        # target_token_count: int = 500, # AppConfigから取得するため削除
        style: str = "会議用要約"
    ) -> DocumentSummary:
        target_token_count = self.config.summarization_target_tokens # AppConfigから取得
        tokens_used_total = 0 # 要約に使用した総トークン数を追跡
        token_count = count_tokens(text, summarizer_ai_client.model_info.name)

        try:
            with Timer("ドキュメント要約"):
                original_length = len(text)

                if token_count <= target_token_count:
                    logger.info("テキストが十分短いため、要約をスキップします")
                    return DocumentSummary(
                        original_length=original_length, summary=text,
                        summary_length=len(text), compression_ratio=1.0, tokens_used=0
                    )

                if token_count > 4000: # 閾値は適宜調整
                    summary_result = await self._summarize_long_document(
                        text, summarizer_ai_client, target_token_count, style
                    )
                    tokens_used_total += summary_result.tokens_used # _summarize_long_document がトークン数を返すようにする
                    return summary_result # DocumentSummaryをそのまま返す
                
                prompt = self._build_summarization_prompt(text, target_token_count, style)
                response = await summarizer_ai_client.request_completion(
                    user_message=prompt, system_message="あなたは専門的な文書要約の専門家です。"
                )
                summary, tokens_used = extract_content_and_tokens(
                    summarizer_ai_client.model_info.provider, response
                )
                tokens_used_total += tokens_used
                summary = summary.strip()
                summary_length = len(summary)
                compression_ratio = summary_length / original_length if original_length > 0 else 0.0
                
                logger.info(f"要約完了: {original_length}文字 → {summary_length}文字 (圧縮率: {compression_ratio:.2%}), トークン: {tokens_used_total}")
                return DocumentSummary(
                    original_length=original_length, summary=summary, summary_length=summary_length,
                    compression_ratio=compression_ratio, tokens_used=tokens_used_total
                )
                
        except Exception as e:
            error_msg = f"要約処理エラー: {str(e)}"
            logger.error(error_msg, exc_info=True)
            # エラー時もDocumentSummaryを返すが、エラー情報を含めるか、例外を再raiseするかは設計次第
            # ここでは空の要約とエラー情報を返すより、例外をraiseする方が呼び出し側で対処しやすい
            raise RuntimeError(error_msg) from e
    
    async def _summarize_long_document(
        self,
        text: str,
        summarizer_ai_client: BaseAIClient,
        target_token_count: int,
        style: str
    ) -> DocumentSummary:
        original_length = len(text)
        tokens_used_total = 0

        chunks = chunk_text(text, max_chunk_size=3000, overlap=200) # chunk_sizeはモデルのコンテキスト長に応じて調整
        logger.info(f"長文書要約開始: {len(chunks)}チャンクに分割")
        
        chunk_summaries = []
        for i, chunk in enumerate(chunks):
            chunk_prompt = self._build_chunk_summarization_prompt(chunk, i + 1, len(chunks))
            response = await summarizer_ai_client.request_completion(
                user_message=chunk_prompt, system_message="あなたは文書要約の専門家です。"
            )
            content, tokens_used = extract_content_and_tokens(
                summarizer_ai_client.model_info.provider, response
            )
            tokens_used_total += tokens_used
            chunk_summaries.append(content.strip())
            logger.debug(f"チャンク {i + 1}/{len(chunks)} 要約完了 (トークン: {tokens_used})")
        
        combined_summaries = "\n\n".join(chunk_summaries)
        final_prompt = self._build_final_summarization_prompt(combined_summaries, target_token_count, style)
        response = await summarizer_ai_client.request_completion(
            user_message=final_prompt, system_message="あなたは文書要約の専門家です。"
        )
        final_summary, tokens_used = extract_content_and_tokens(
            summarizer_ai_client.model_info.provider, response
        )
        tokens_used_total += tokens_used

        final_summary = final_summary.strip()
        summary_length = len(final_summary)
        compression_ratio = summary_length / original_length if original_length > 0 else 0.0
        
        logger.info(f"段階的要約完了: {original_length}文字 → {summary_length}文字 (圧縮率: {compression_ratio:.2%}), 総トークン: {tokens_used_total}")
        return DocumentSummary(
            original_length=original_length, summary=final_summary, summary_length=summary_length,
            compression_ratio=compression_ratio, tokens_used=tokens_used_total
        )
    
    def _build_summarization_prompt(self, text: str, target_token_count: int, style: str) -> str:
        return f"""以下のドキュメントを{style}として、約{target_token_count}トークン（日本語で約{int(target_token_count * 1.5)}～{target_token_count * 2}文字）以内で要約してください。

要約の指針：
1. 主要なポイントと重要な詳細を含める
2. 論理的な構造を保持する
3. 会議での議論に必要な情報を優先する
4. 専門用語は適切に説明するか、文脈から理解できるようにする

【要約対象ドキュメント】
{text}

【要約結果】"""
    
    def _build_chunk_summarization_prompt(self, chunk: str, chunk_num: int, total_chunks: int) -> str:
        return f"""以下は長いドキュメントの一部（{chunk_num}/{total_chunks}）です。この部分の主要な情報を保持しつつ、できるだけ簡潔に300～500文字程度で要約してください。後でこれらの部分要約を結合して最終的な要約を作成します。

【ドキュメント（部分 {chunk_num}/{total_chunks}）】
{chunk}

【この部分の要約】"""
    
    def _build_final_summarization_prompt(self, combined_summaries: str, target_token_count: int, style: str) -> str:
        return f"""以下は長いドキュメントを部分ごとに要約したものです。これらを統合し、重複を避け、全体の流れと論理性を重視して、約{target_token_count}トークン（日本語で約{int(target_token_count * 1.5)}～{target_token_count * 2}文字）以内の首尾一貫した{style}を作成してください。

要約の指針：
1. 各部分の要点を統合し、全体の流れを重視する
2. 重複を避け、一貫性のある要約にする
3. 会議での議論に必要な情報を優先する
4. 論理的な構造で整理する

【部分ごとの要約】
{combined_summaries}

【最終要約】"""

# ヘルパー関数は DocumentProcessor のインスタンス化方法が変わったため、修正が必要
# 現状、main.pyからは直接使われていないため、一旦コメントアウトまたは削除を検討。
# もし外部から使う場合は、configを渡すように修正が必要。
# def extract_document_text(file_path: str, config: AppConfig) -> ExtractionResult:
#     processor = DocumentProcessor(config)
#     return processor.extract_text(file_path)

# async def summarize_document(
#     text: str,
#     ai_client: BaseAIClient,
#     config: AppConfig
#     # target_tokens は config から取得
# ) -> DocumentSummary:
#     processor = DocumentProcessor(config)
#     return await processor.summarize_document_for_meeting(text, ai_client)